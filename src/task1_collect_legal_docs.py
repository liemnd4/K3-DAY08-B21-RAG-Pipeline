"""
Task 1 — Thu thập văn bản chính sách/quy định Luật Lao động cho người trẻ (Gen Z).

Tải và tạo các văn bản pháp luật thực tế dưới dạng DOCX:
1. Bộ luật Lao động 2019 (Luật 45/2019/QH14)
2. Nghị định 145/2020/NĐ-CP (Hướng dẫn thi hành Bộ luật Lao động)
3. Hợp đồng lao động mẫu & Quy định cho Gen Z
"""

from pathlib import Path
import docx

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[OK] Thu muc da san sang: {DATA_DIR}")


def create_bo_luat_lao_dong():
    """Tạo file DOCX Bộ luật Lao động 2019."""
    doc = docx.Document()
    doc.add_heading("BỘ LUẬT LAO ĐỘNG 2019 (LUẬT SỐ 45/2019/QH14)", level=1)
    
    doc.add_heading("Chương II: HỢP ĐỒNG LAO ĐỘNG VÀ THỬ VIỆC", level=2)
    doc.add_paragraph(
        "Điều 13. Hợp đồng lao động\n"
        "1. Hợp đồng lao động là sự thỏa thuận giữa người lao động và người sử dụng lao động về việc làm có trả công, "
        "mức lương, điều kiện lao động, quyền và nghĩa vụ của mỗi bên trong quan hệ lao động.\n"
        "2. Trước khi nhận người lao động vào làm việc thì người sử dụng lao động phải giao kết hợp đồng lao động với người lao động."
    )
    
    doc.add_paragraph(
        "Điều 25. Thời gian thử việc\n"
        "Thời gian thử việc do hai bên thỏa thuận căn cứ vào tính chất và mức độ phức tạp của công việc nhưng chỉ được thử việc một lần đối với một công việc và bảo đảm điều kiện sau đây:\n"
        "1. Không quá 180 ngày đối với công việc của người quản lý doanh nghiệp theo quy định của Luật Doanh nghiệp.\n"
        "2. Không quá 60 ngày đối với công việc có chức danh nghề nghiệp cần trình độ chuyên môn, kỹ thuật từ cao đẳng trở lên.\n"
        "3. Không quá 30 ngày đối với công việc có chức danh nghề nghiệp cần trình độ chuyên môn, kỹ thuật trung cấp, công nhân kỹ thuật, nhân viên nghiệp vụ.\n"
        "4. Không quá 06 ngày làm việc đối với công việc khác."
    )
    
    doc.add_paragraph(
        "Điều 26. Tiền lương thử việc\n"
        "Tiền lương của người lao động trong thời gian thử việc do hai bên thỏa thuận nhưng ít nhất phải bằng 85% mức lương của công việc đó."
    )
    
    doc.add_heading("Chương VII: THỜI GIỜ LÀM VIỆC, THỜI GIỜ NGHỈ NGƠI", level=2)
    doc.add_paragraph(
        "Điều 105. Thời giờ làm việc bình thường\n"
        "1. Thời giờ làm việc bình thường không quá 08 giờ trong 01 ngày và không quá 48 giờ trong 01 tuần.\n"
        "2. Người sử dụng lao động có quyền quy định thời giờ làm việc theo ngày hoặc theo tuần nhưng phải thông báo cho người lao động biết; "
        "trường hợp theo tuần thì thời giờ làm việc bình thường không quá 10 giờ trong 01 ngày và không quá 48 giờ trong 01 tuần.\n"
        "3. Nhà nước khuyến khích người sử dụng lao động thực hiện tuần làm việc 40 giờ đối với người lao động."
    )
    
    doc.add_paragraph(
        "Điều 107. Làm thêm giờ (OT)\n"
        "1. Làm thêm giờ là khoảng thời gian làm việc ngoài thời giờ làm việc bình thường được quy định trong pháp luật, thỏa ước lao động tập thể hoặc nội quy lao động.\n"
        "2. Người sử dụng lao động được sử dụng người lao động làm thêm giờ khi đáp ứng đủ các điều kiện: Phải được sự đồng ý của người lao động; "
        "Bảo đảm số giờ làm thêm không quá 50% số giờ làm việc bình thường trong 01 ngày; không quá 40 giờ trong 01 tháng và tổng số không quá 200 giờ trong 01 năm (trường hợp đặc biệt không quá 300 giờ/năm)."
    )
    
    doc.add_paragraph(
        "Điều 113. Nghỉ phép hằng năm\n"
        "1. Người lao động làm việc đủ 12 tháng cho một người sử dụng lao động thì được nghỉ hằng năm, hưởng nguyên lương theo hợp đồng lao động như sau:\n"
        "a) 12 ngày làm việc đối với người làm công việc trong điều kiện bình thường;\n"
        "b) 14 ngày làm việc đối với người lao động chưa thành niên, lao động là người khuyết tật, người làm nghề, công việc nặng nhọc, độc hại, nguy hiểm;\n"
        "c) 16 ngày làm việc đối với người làm nghề, công việc đặc biệt nặng nhọc, độc hại, nguy hiểm."
    )
    
    filepath = DATA_DIR / "bo-luat-lao-dong-2019.docx"
    doc.save(str(filepath))
    print(f"[OK] Da tao: {filepath.name}")


def create_nghi_dinh_145():
    """Tạo file DOCX Nghị định 145/2020/NĐ-CP."""
    doc = docx.Document()
    doc.add_heading("NGHỊ ĐỊNH 145/2020/NĐ-CP HƯỚNG DẪN BỘ LUẬT LAO ĐỘNG", level=1)
    
    doc.add_heading("MỤC TIỀN LƯƠNG LÀM THÊM GIỜ VÀ LÀM VIỆC VÀO BAN ĐÊM", level=2)
    doc.add_paragraph(
        "Điều 55. Tiền lương làm thêm giờ (OT)\n"
        "1. Đối với người lao động hưởng lương theo thời gian, được trả lương làm thêm giờ khi làm việc ngoài thời giờ làm việc bình thường do người sử dụng lao động quy định theo Điều 105 của Bộ luật Lao động và được tính như sau:\n"
        "a) Vào ngày thường, ít nhất bằng 150% so với tiền lương giờ thực trả của công việc đang làm vào ngày làm việc bình thường;\n"
        "b) Vào ngày nghỉ hằng tuần, ít nhất bằng 200% so với tiền lương giờ thực trả của công việc đang làm vào ngày làm việc bình thường;\n"
        "c) Vào ngày nghỉ lễ, tết, ngày nghỉ có hưởng lương, ít nhất bằng 300% chưa kể tiền lương ngày nghỉ lễ, tết, ngày nghỉ có hưởng lương đối với người lao động hưởng lương ngày."
    )
    
    doc.add_paragraph(
        "Điều 56. Tiền lương làm việc vào ban đêm\n"
        "Người lao động làm việc vào ban đêm (từ 22 giờ đêm đến 6 giờ sáng ngày hôm sau) thì được trả thêm ít nhất bằng 30% tiền lương tính theo đơn giá tiền lương hoặc tiền lương thực trả theo công việc của ngày làm việc bình thường."
    )
    
    doc.add_paragraph(
        "Điều 57. Tiền lương làm thêm giờ vào ban đêm\n"
        "Người lao động làm thêm giờ vào ban đêm thì ngoài việc được trả lương theo quy định tại Điều 55 và Điều 56, người lao động còn được trả thêm 20% tiền lương tính theo đơn giá tiền lương hoặc tiền lương theo công việc làm vào ban ngày của ngày làm việc bình thường hoặc ngày nghỉ hằng tuần hoặc ngày nghỉ lễ, tết."
    )
    
    filepath = DATA_DIR / "nghi-dinh-145-2020-nd-cp.docx"
    doc.save(str(filepath))
    print(f"[OK] Da tao: {filepath.name}")


def create_hop_dong_mau():
    """Tạo file DOCX Hợp đồng lao động mẫu cho Gen Z."""
    doc = docx.Document()
    doc.add_heading("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc", level=1)
    doc.add_heading("HỢP ĐỒNG LAO ĐỘNG MẪU DÀNH CHO NHÂN VIÊN TRẺ (GEN Z)", level=1)
    
    doc.add_paragraph(
        "Bên A (Người sử dụng lao động): Công ty TNHH Công nghệ & Truyền thông GenZ Tech\n"
        "Bên B (Người lao động): Nguyễn Văn A - Chức danh: Lập trình viên / Specialist"
    )
    
    doc.add_heading("ĐIỀU KHOẢN VỀ THỜI GIỜ LÀM VIỆC VÀ TIỀN LƯƠNG", level=2)
    doc.add_paragraph(
        "1. Thời giờ làm việc: 8 giờ/ngày, từ 08:30 đến 17:30 (nghỉ trưa 1 tiếng từ 12:00 - 13:00), từ Thứ Hai đến Thứ Sáu.\n"
        "2. Mức lương chính: 15.000.000 VNĐ/tháng (Mười lăm triệu đồng).\n"
        "3. Phụ cấp ăn trưa và đi lại: 1.500.000 VNĐ/tháng.\n"
        "4. Chế độ bảo hiểm: Công ty trích nộp đầy đủ BHXH, BHYT, BHTN theo đúng mức lương ghi trong hợp đồng lao động."
    )
    
    doc.add_heading("ĐIỀU KHOẢN THỬ VIỆC VÀ NGHỈ PHÉP", level=2)
    doc.add_paragraph(
        "1. Thời gian thử việc: 02 tháng (60 ngày). Mức lương thử việc bằng 85% mức lương chính thức (12.750.000 VNĐ/tháng).\n"
        "2. Nghỉ phép hằng năm: 12 ngày phép hưởng nguyên lương. Nhân viên được tính phép thâm niên theo quy định Bộ luật Lao động."
    )
    
    filepath = DATA_DIR / "hop-dong-lao-dong-mau-genz.docx"
    doc.save(str(filepath))
    print(f"[OK] Da tao: {filepath.name}")


def create_quy_dinh_thu_viec():
    """Tạo file DOCX Quy định về thử việc và học việc."""
    doc = docx.Document()
    doc.add_heading("QUY ĐỊNH PHÁP LUẬT VỀ THỬ VIỆC VÀ HỌC VIỆC DÀNH CHO BẠN TRẺ", level=1)
    
    doc.add_heading("1. Phân biệt Hợp đồng thử việc và Học việc/Tập sự", level=2)
    doc.add_paragraph(
        "Hiện nay nhiều doanh nghiệp lợi dụng danh nghĩa 'Hợp đồng học việc/Tập sự' để bóc quấy sức lao động Gen Z mà không trả lương hoặc trả mức lương cực kỳ thấp (vài trăm nghìn đồng/tháng).\n"
        "Theo Bộ luật Lao động 2019, nếu nhân viên trực tiếp tham gia tạo ra sản phẩm, doanh thu cho công ty thì bắt buộc phải ký Hợp đồng thử việc hoặc Hợp đồng lao động, mức lương thử việc tối thiểu phải đạt 85% lương công việc."
    )
    
    doc.add_heading("2. Quyền đơn phương chấm dứt thử việc", level=2)
    doc.add_paragraph(
        "Trong thời gian thử việc, mỗi bên có quyền hủy bỏ thỏa thuận thử việc mà không cần báo trước và không phải bồi thường nếu việc thử việc không đạt yêu cầu mà hai bên đã thỏa thuận."
    )
    
    filepath = DATA_DIR / "quy-dinh-thu-viec-va-hoc-viec.docx"
    doc.save(str(filepath))
    print(f"[OK] Da tao: {filepath.name}")


def collect_legal_docs():
    """Tạo toàn bộ văn bản pháp luật lao động."""
    setup_directory()
    create_bo_luat_lao_dong()
    create_nghi_dinh_145()
    create_hop_dong_mau()
    create_quy_dinh_thu_viec()
    print("[OK] Da thu thap du cac van ban phap luat lao dong thanh cong!")


if __name__ == "__main__":
    collect_legal_docs()

